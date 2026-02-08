"""
Generate all DOCX documents for the ASPR Photo Repository project.
Reads markdown source files from docs/ and produces branded DOCX output.

Run:  python scripts/generate_all_docx.py
Requires: pip install python-docx
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"

# ── ASPR / HHS brand colours ──────────────────────────────────────────
BLUE_DARK      = RGBColor(0x06, 0x2E, 0x61)
BLUE_PRIMARY   = RGBColor(0x15, 0x51, 0x97)
GOLD           = RGBColor(0xAA, 0x64, 0x04)
RED            = RGBColor(0x99, 0x00, 0x00)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY     = RGBColor(0xF2, 0xF2, 0xF2)

BLUE_DARK_HEX    = "062E61"
BLUE_PRIMARY_HEX = "155197"
GOLD_HEX         = "AA6404"
LIGHT_GRAY_HEX   = "F2F2F2"

# ── Logo Paths ────────────────────────────────────────────────────────
ASPR_LOGO = ROOT / "public" / "aspr-logo-blue.png"
LEIDOS_LOGO = Path(
    r"C:\Users\ravinder.mathaudhu\OneDrive - HHS Office of the Secretary"
    r"\Documents\Projects\New folder\Leidos-Logo-Suite\Leidos-Logo-Suite"
    r"\02-Digital\03-Raster-PNG\Leidos-logo-horz-full-rgb-@2x.png"
)


# ══════════════════════════════════════════════════════════════════════
#  DOCX HELPERS (shared across all documents)
# ══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, hex_color):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def styled_table(doc, headers, rows, col_widths=None):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, BLUE_DARK_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, text in enumerate(row_data[:ncols]):
            cell = row.cells[ci]
            if ri % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY_HEX)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    if col_widths:
        total = sum(col_widths)
        table_width = Inches(6.5)
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = int(table_width * col_widths[ci] / total)

    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    color_map = {1: BLUE_DARK, 2: BLUE_PRIMARY, 3: GOLD}
    color = color_map.get(level, BLUE_DARK)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11),
             color=None, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_toc(doc):
    """Add a Word field-based Table of Contents."""
    add_heading_styled(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()

    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run("[Right-click → Update Field to generate Table of Contents]")
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run4.font.size = Pt(10)
    run4.italic = True

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


def setup_doc(doc_title, doc_subtitle, version="1.0", date="February 7, 2026",
              status="Draft"):
    """Create a new Document with branding, cover page, and TOC."""
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"ASPR Photo Repository — {doc_title}")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = BLUE_PRIMARY
    run.font.name = "Calibri"

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("HHS/ASPR — For Official Use Only | Leidos")
    run.font.size = Pt(8)
    run.font.color.rgb = BLUE_DARK
    run.font.name = "Calibri"

    # ── Cover Page ──
    for _ in range(3):
        doc.add_paragraph()

    # Logos side by side
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ASPR_LOGO.exists():
        run = logo_para.add_run()
        run.add_picture(str(ASPR_LOGO), width=Inches(2.0))
    if LEIDOS_LOGO.exists():
        run = logo_para.add_run("     ")  # spacer
        run = logo_para.add_run()
        run.add_picture(str(LEIDOS_LOGO), width=Inches(2.0))

    doc.add_paragraph()

    add_para(doc, "U.S. Department of Health and Human Services",
             size=Pt(12), color=BLUE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Administration for Strategic Preparedness and Response (ASPR)",
             size=Pt(11), color=BLUE_PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(2):
        doc.add_paragraph()

    add_para(doc, doc_title,
             bold=True, size=Pt(26), color=BLUE_DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_para(doc, doc_subtitle,
             size=Pt(18), color=BLUE_PRIMARY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

    for _ in range(2):
        doc.add_paragraph()

    # Document info table
    styled_table(doc,
        ["Property", "Value"],
        [
            ["Document Version", version],
            ["Date", date],
            ["Application Version", "0.1.0"],
            ["Project", "app-aspr-photos-lab"],
            ["Status", status],
            ["Classification", "For Official Use Only"],
            ["Federal Project Sponsor", "[Name]"],
        ],
        col_widths=[35, 65],
    )

    for _ in range(2):
        doc.add_paragraph()

    add_para(doc, "DRAFT — FOR REVIEW",
             bold=True, size=Pt(14), color=RED,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # TOC
    add_toc(doc)
    doc.add_page_break()

    return doc


# ══════════════════════════════════════════════════════════════════════
#  MARKDOWN → DOCX CONVERTER
# ══════════════════════════════════════════════════════════════════════

def parse_md_table(lines):
    """Parse markdown table lines into (headers, rows)."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if i == 0:
            headers = cells
        elif i == 1:
            continue  # separator row
        else:
            rows.append(cells)
    return headers, rows


def md_to_docx(md_path, doc_title, doc_subtitle, out_filename):
    """Convert a markdown file to a branded DOCX document."""
    md_text = md_path.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = setup_doc(doc_title, doc_subtitle)

    # Skip the markdown header block (title, metadata table, ---)
    # Find where the actual content starts (first ## heading)
    content_start = 0
    found_first_heading = False
    for i, line in enumerate(lines):
        if line.startswith('## ') and not found_first_heading:
            # Skip "Table of Contents" if present
            if 'table of contents' in line.lower():
                continue
            content_start = i
            found_first_heading = True
            break

    i = content_start
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip markdown TOC links
        if stripped.startswith('- [') and '](#' in stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            if text.startswith('#'):
                text = text.lstrip('#').strip()
            add_heading_styled(doc, text, level=3)
            i += 1
            continue

        if stripped.startswith('## '):
            text = stripped[3:].strip()
            add_heading_styled(doc, text, level=2)
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Table detection
        if '|' in stripped and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            if headers and rows:
                styled_table(doc, headers, rows)
            continue

        # Code block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(code_text)
            run.font.size = Pt(8.5)
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Bullet point
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            # Remove bold markers
            text = text.replace('**', '')
            level = 1 if line.startswith('  ') else 0
            add_bullet(doc, text, level=level)
            i += 1
            continue

        # Bold paragraph
        if stripped.startswith('**') and stripped.endswith('**'):
            text = stripped.strip('*').strip()
            add_para(doc, text, bold=True)
            i += 1
            continue

        # Regular paragraph
        text = stripped.replace('**', '')
        if text:
            add_para(doc, text)
        i += 1

    # Save
    out_path = DOCS / out_filename
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size / 1024
    print(f"  [OK] {out_filename} ({size_kb:.1f} KB)")
    return out_path


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT DEFINITIONS
# ══════════════════════════════════════════════════════════════════════

DOCUMENTS = [
    {
        "md": "01_SRS_Software_Requirements_Specification.md",
        "title": "Software Requirements Specification",
        "subtitle": "ASPR Photo Repository Application",
        "out": "01_ASPR_Photos_SRS.docx",
    },
    {
        "md": "02_SDD_System_Design_Document.md",
        "title": "System Design Document",
        "subtitle": "ASPR Photo Repository Application",
        "out": "02_ASPR_Photos_SDD.docx",
    },
    {
        "md": "03_Security_Plan.md",
        "title": "Security Plan",
        "subtitle": "ASPR Photo Repository Application",
        "out": "03_ASPR_Photos_Security_Plan.docx",
    },
    {
        "md": "04_Deployment_Operations_Guide.md",
        "title": "Deployment & Operations Guide",
        "subtitle": "ASPR Photo Repository Application",
        "out": "04_ASPR_Photos_Deployment_Guide.docx",
    },
    {
        "md": "05_User_Guide.md",
        "title": "User Guide",
        "subtitle": "ASPR Photo Repository Application",
        "out": "05_ASPR_Photos_User_Guide.docx",
    },
    {
        "md": "06_API_Data_Reference.md",
        "title": "API & Data Reference",
        "subtitle": "ASPR Photo Repository Application",
        "out": "06_ASPR_Photos_API_Reference.docx",
    },
]


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("  ASPR Photo Repository — Document Generation")
    print("=" * 60)
    print()

    # Check logos
    print(f"  ASPR Logo:   {'[OK] Found' if ASPR_LOGO.exists() else '[!] Not found'}")
    print(f"  Leidos Logo: {'[OK] Found' if LEIDOS_LOGO.exists() else '[!] Not found'}")
    print()

    generated = []
    errors = []

    for doc_def in DOCUMENTS:
        md_path = DOCS / doc_def["md"]
        if not md_path.exists():
            print(f"  [!] Skipping {doc_def['md']} (not found)")
            errors.append(doc_def["md"])
            continue

        try:
            out_path = md_to_docx(
                md_path,
                doc_def["title"],
                doc_def["subtitle"],
                doc_def["out"],
            )
            generated.append(out_path)
        except Exception as e:
            print(f"  [ERR] Error generating {doc_def['out']}: {e}")
            errors.append(doc_def["out"])

    print()
    print(f"  Generated: {len(generated)} documents")
    if errors:
        print(f"  Errors:    {len(errors)} — {', '.join(errors)}")
    print()
    print("  Done! Open documents in Word and right-click TOC > Update Field")
    print("=" * 60)
